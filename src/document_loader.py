"""
文件載入器 - 處理不同格式的文件
"""
import os
from typing import List, Dict
from pathlib import Path

class Document:
    """文件物件"""
    def __init__(self, content: str, metadata: Dict = None):
        self.content = content
        self.metadata = metadata or {}
    
    def __repr__(self):
        return f"Document(content={self.content[:50]}..., metadata={self.metadata})"


class DocumentLoader:
    """文件載入器 - 支援多種格式"""
    
    def __init__(self, data_path: str):
        self.data_path = data_path
    
    def load_txt(self, file_path: str) -> List[Document]:
        """載入文字檔"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 分段處理（以空行分隔）
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        documents = []
        for i, paragraph in enumerate(paragraphs):
            doc = Document(
                content=paragraph,
                metadata={
                    'source': file_path,
                    'paragraph': i,
                    'type': 'txt'
                }
            )
            documents.append(doc)
        
        return documents
    
    def load_pdf(self, file_path: str) -> List[Document]:
        """載入 PDF 檔案"""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(file_path)
            documents = []
            
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    doc = Document(
                        content=text,
                        metadata={
                            'source': file_path,
                            'page': page_num + 1,
                            'type': 'pdf'
                        }
                    )
                    documents.append(doc)
            
            return documents
        except ImportError:
            print("請安裝 pypdf: pip install pypdf")
            return []
    
    def load_docx(self, file_path: str) -> List[Document]:
        """載入 Word 檔案"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            documents = []
            
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    document = Document(
                        content=para.text,
                        metadata={
                            'source': file_path,
                            'paragraph': i,
                            'type': 'docx'
                        }
                    )
                    documents.append(document)
            
            return documents
        except ImportError:
            print("請安裝 python-docx: pip install python-docx")
            return []
    
    def load_directory(self, directory: str = None) -> List[Document]:
        """載入整個資料夾的文件"""
        if directory is None:
            directory = self.data_path
        
        documents = []
        path = Path(directory)
        
        if not path.exists():
            print(f"資料夾不存在: {directory}")
            return documents
        
        # 支援的檔案格式
        for file_path in path.rglob('*'):
            if file_path.is_file():
                file_ext = file_path.suffix.lower()
                
                try:
                    if file_ext == '.txt':
                        docs = self.load_txt(str(file_path))
                        documents.extend(docs)
                        print(f"✓ 載入 {file_path.name}: {len(docs)} 個文件片段")
                    
                    elif file_ext == '.pdf':
                        docs = self.load_pdf(str(file_path))
                        documents.extend(docs)
                        print(f"✓ 載入 {file_path.name}: {len(docs)} 頁")
                    
                    elif file_ext in ['.docx', '.doc']:
                        docs = self.load_docx(str(file_path))
                        documents.extend(docs)
                        print(f"✓ 載入 {file_path.name}: {len(docs)} 個段落")
                
                except Exception as e:
                    print(f"✗ 載入 {file_path.name} 失敗: {e}")
        
        print(f"\n總共載入 {len(documents)} 個文件片段")
        return documents
    
    def chunk_documents(self, documents: List[Document], chunk_size: int = 500, overlap: int = 50) -> List[Document]:
        """將文件分割成更小的區塊"""
        chunked_docs = []
        
        for doc in documents:
            content = doc.content
            
            # 如果文件已經夠小，直接使用
            if len(content) <= chunk_size:
                chunked_docs.append(doc)
                continue
            
            # 分割成區塊
            for i in range(0, len(content), chunk_size - overlap):
                chunk = content[i:i + chunk_size]
                if chunk.strip():
                    new_doc = Document(
                        content=chunk,
                        metadata={
                            **doc.metadata,
                            'chunk': i // (chunk_size - overlap)
                        }
                    )
                    chunked_docs.append(new_doc)
        
        return chunked_docs
